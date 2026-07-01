"""Word (.docx) parser: paragraphs as text, tables preserved.

Word has no reliable page model without rendering, so the whole document is
treated as a single logical page. Paragraph order and tables are preserved.
"""

from __future__ import annotations

import io

from audit_assistant.core.exceptions import ParsingError
from audit_assistant.core.logging import get_logger
from audit_assistant.domain.models import DocumentPage, FileType, ParsedDocument, Table

log = get_logger(__name__)


class DocxParser:
    """Parses .docx documents via python-docx."""

    def supports(self, file_type: FileType) -> bool:
        return file_type == FileType.DOCX

    def parse(self, *, filename: str, data: bytes, file_type: FileType) -> ParsedDocument:
        import docx

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001
            raise ParsingError(f"Failed to parse Word doc '{filename}': {exc}") from exc

        paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paragraphs)

        tables: list[Table] = []
        for tbl in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            if rows:
                tables.append(Table(rows=rows, page=1))

        core = document.core_properties
        metadata = {
            k: str(v)
            for k, v in {
                "title": core.title,
                "author": core.author,
                "created": core.created,
            }.items()
            if v
        }

        log.info("Parsed Word doc '%s': %d paragraphs, %d tables", filename, len(paragraphs), len(tables))
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            pages=[DocumentPage(number=1, text=text, tables=tables)],
            metadata=metadata,
        )
