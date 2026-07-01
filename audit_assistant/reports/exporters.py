"""Report exporters: Report -> PDF / Word / Excel bytes (English + Arabic/RTL)."""

from __future__ import annotations

import io

from audit_assistant.reports.models import Report
from audit_assistant.reports.rtl import ARABIC_FONT_PATH, contains_arabic, shape_for_pdf


def _latin1_safe(text: str) -> str:
    """fpdf2 core fonts are Latin-1 only; replace unsupported glyphs."""
    return text.encode("latin-1", "replace").decode("latin-1")


def to_pdf(report: Report) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    has_arabic_font = ARABIC_FONT_PATH.exists()
    if has_arabic_font:
        pdf.add_font("amiri", "", str(ARABIC_FONT_PATH))

    def block(text: str, size: int, style: str = "") -> None:
        # Arabic blocks: Amiri font, shaped glyphs, right-aligned. Others: Helvetica.
        if has_arabic_font and contains_arabic(text):
            pdf.set_font("amiri", "", size)
            pdf.multi_cell(0, size * 0.6 + 2, shape_for_pdf(text), align="R",
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.set_font("Helvetica", style, size)
            pdf.multi_cell(0, size * 0.5 + 2, _latin1_safe(text),
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    block(report.title, 18, "B")
    if report.subtitle:
        block(report.subtitle, 12, "I")
    block(f"{report.author} - {report.report_date}", 10)
    pdf.ln(4)

    for section in report.sections:
        block(section.heading, 13, "B")
        block(section.body, 11)
        pdf.ln(3)

    return bytes(pdf.output())


def _set_paragraph_rtl(paragraph) -> None:
    """Mark a python-docx paragraph as right-to-left (Arabic)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:bidi"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        r_pr.append(OxmlElement("w:rtl"))


def to_docx(report: Report) -> bytes:
    import docx

    document = docx.Document()

    def add(text: str, adder) -> None:
        para = adder(text)
        if contains_arabic(text):
            _set_paragraph_rtl(para)

    add(report.title, lambda t: document.add_heading(t, level=0))
    if report.subtitle:
        add(report.subtitle, lambda t: document.add_paragraph(t, style="Intense Quote"))
    document.add_paragraph(f"{report.author} — {report.report_date}")

    for section in report.sections:
        add(section.heading, lambda t: document.add_heading(t, level=1))
        for para in section.body.split("\n\n"):
            if para.strip():
                add(para.strip(), document.add_paragraph)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def to_xlsx(report: Report) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Report"

    report_has_arabic = contains_arabic(report.to_markdown())
    if report_has_arabic:
        ws.sheet_view.rightToLeft = True
    align = Alignment(wrap_text=True, vertical="top",
                      horizontal="right" if report_has_arabic else "left")

    ws["A1"] = report.title
    ws["A1"].font = Font(bold=True, size=16)
    ws["A2"] = report.subtitle
    ws["A3"] = f"{report.author} — {report.report_date}"

    row = 5
    for section in report.sections:
        cell = ws.cell(row=row, column=1, value=section.heading)
        cell.font = Font(bold=True, size=12)
        row += 1
        body_cell = ws.cell(row=row, column=1, value=section.body)
        body_cell.alignment = align
        row += 2

    ws.column_dimensions[get_column_letter(1)].width = 100
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


EXPORTERS = {"PDF": to_pdf, "Word": to_docx, "Excel": to_xlsx}
EXTENSIONS = {"PDF": "pdf", "Word": "docx", "Excel": "xlsx"}
MIME_TYPES = {
    "PDF": "application/pdf",
    "Word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "Excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
