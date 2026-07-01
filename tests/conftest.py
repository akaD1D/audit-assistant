"""Shared test fixtures: real, in-memory sample files for each parser."""

from __future__ import annotations

import io

import pytest


@pytest.fixture
def pdf_bytes() -> bytes:
    import fitz  # PyMuPDF

    doc = fitz.open()
    for i in range(2):
        page = doc.new_page()
        page.insert_text((72, 72), f"Materiality threshold on page {i + 1}. Revenue recognition.")
    data = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def xlsx_bytes() -> bytes:
    import pandas as pd

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        pd.DataFrame({"Account": ["Cash", "AR"], "Amount": [1000, 500]}).to_excel(
            writer, sheet_name="TrialBalance", index=False
        )
        pd.DataFrame({"Vendor": ["Acme"], "Total": [250]}).to_excel(
            writer, sheet_name="Payables", index=False
        )
    return buf.getvalue()


@pytest.fixture
def csv_bytes() -> bytes:
    return b"Account,Amount\nCash,1000\nAR,500\n"


@pytest.fixture
def docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_paragraph("Revenue recognition policy under IFRS 15.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Account"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Cash"
    table.cell(1, 1).text = "1000"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


@pytest.fixture
def txt_bytes() -> bytes:
    return "ISA 315 risk assessment procedures.".encode("utf-8")


@pytest.fixture
def png_bytes() -> bytes:
    from PIL import Image

    img = Image.new("RGB", (80, 40), (255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()
